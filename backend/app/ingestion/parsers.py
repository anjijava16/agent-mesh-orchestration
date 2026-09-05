"""File parsing.

Returns page-aware text so citations can carry a page number. Every parser
degrades to plain-text decoding rather than raising, because a partially
readable document is more useful than a failed job.
"""
from __future__ import annotations

import csv
import io
import json
from dataclasses import dataclass

from app.core.logging import get_logger

log = get_logger(__name__)


@dataclass
class ParsedPage:
    page: int
    text: str


def parse(data: bytes, filename: str, content_type: str) -> list[ParsedPage]:
    lower = filename.lower()
    try:
        if lower.endswith(".pdf") or content_type == "application/pdf":
            return _pdf(data)
        if lower.endswith((".docx", ".doc")):
            return _docx(data)
        if lower.endswith((".xlsx", ".xls")):
            return _xlsx(data)
        if lower.endswith(".csv"):
            return _csv(data)
        if lower.endswith((".html", ".htm")):
            return _html(data)
        if lower.endswith(".json"):
            return [ParsedPage(1, json.dumps(json.loads(data.decode("utf-8", "replace")), indent=2))]
    except Exception as exc:
        log.warning("parser_fallback", filename=filename, error=str(exc)[:200])
    return [ParsedPage(1, data.decode("utf-8", errors="replace"))]


def _pdf(data: bytes) -> list[ParsedPage]:
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(data))
    pages = [ParsedPage(i, (p.extract_text() or "").strip()) for i, p in enumerate(reader.pages, start=1)]
    if any(p.text for p in pages):
        return [p for p in pages if p.text]
    # No embedded text - a scan. Hand it to OCR if it is available.
    return _ocr(data)


def _ocr(data: bytes) -> list[ParsedPage]:
    try:
        import pytesseract
        from pdf2image import convert_from_bytes

        images = convert_from_bytes(data, dpi=200)
        return [ParsedPage(i, pytesseract.image_to_string(img)) for i, img in enumerate(images, start=1)]
    except Exception as exc:
        log.warning("ocr_unavailable", error=str(exc)[:200])
        return [ParsedPage(1, "")]


def _docx(data: bytes) -> list[ParsedPage]:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
    return [ParsedPage(1, "\n".join(paragraphs))]


def _xlsx(data: bytes) -> list[ParsedPage]:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    pages = []
    for index, sheet in enumerate(workbook.worksheets, start=1):
        rows = [
            " | ".join("" if c is None else str(c) for c in row)
            for row in sheet.iter_rows(values_only=True)
        ]
        pages.append(ParsedPage(index, f"# Sheet: {sheet.title}\n" + "\n".join(rows)))
    return pages


def _csv(data: bytes) -> list[ParsedPage]:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    return [ParsedPage(1, "\n".join(" | ".join(row) for row in reader))]


def _html(data: bytes) -> list[ParsedPage]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data.decode("utf-8", errors="replace"), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return [ParsedPage(1, soup.get_text("\n", strip=True))]
